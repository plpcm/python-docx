# encoding: utf-8

"""
Step implementations for graphical object (shape) related features
"""

from __future__ import absolute_import, print_function, unicode_literals

import hashlib

from behave import given, then, when

from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE
from docx.parts.document import InlineShape, InlineShapes

from .helpers import test_docx, test_file_path


# given ===================================================

@given('a document containing five inline shapes')
def given_a_document_containing_two_inline_shapes(context):
    docx_path = test_docx('shp-inline-shape-access')
    context.document = Document(docx_path)


@given('an inline shape collection containing five shapes')
def given_inline_shape_collection_containing_two_shapes(context):
    docx_path = test_docx('shp-inline-shape-access')
    document = Document(docx_path)
    context.inline_shapes = document.inline_shapes


@given('an inline shape known to be {shp_of_type}')
def given_inline_shape_known_to_be_shape_of_type(context, shp_of_type):
    inline_shape_idx = {
        'an embedded picture':  0,
        'a linked picture':     1,
        'a link+embed picture': 2,
        'a smart art diagram':  3,
        'a chart':              4,
    }[shp_of_type]
    docx_path = test_docx('shp-inline-shape-access')
    document = Document(docx_path)
    context.inline_shape = document.inline_shapes[inline_shape_idx]


# when =====================================================

@when('I add an inline picture from a file-like object')
def when_add_inline_picture_from_file_like_object(context):
    document = context.document
    with open(test_file_path('monty-truth.png')) as f:
        context.inline_shape = (
            document.add_inline_picture(f)
        )


@when('I add an inline picture to the document')
def when_add_inline_picture_to_document(context):
    document = context.document
    context.inline_shape = (
        document.add_inline_picture(test_file_path('monty-truth.png'))
    )


# then =====================================================

@then('I can access each inline shape by index')
def then_can_access_each_inline_shape_by_index(context):
    inline_shapes = context.inline_shapes
    for idx in range(2):
        inline_shape = inline_shapes[idx]
        assert isinstance(inline_shape, InlineShape)


@then('I can access the inline shape collection of the document')
def then_can_access_inline_shape_collection_of_document(context):
    document = context.document
    inline_shapes = document.inline_shapes
    assert isinstance(inline_shapes, InlineShapes)


@then('I can iterate over the inline shape collection')
def then_can_iterate_over_inline_shape_collection(context):
    inline_shapes = context.inline_shapes
    shape_count = 0
    for inline_shape in inline_shapes:
        shape_count += 1
        assert isinstance(inline_shape, InlineShape)
    expected_count = 5
    assert shape_count == expected_count, (
        'expected %d, got %d' % (expected_count, shape_count)
    )


@then('its inline shape type is {shape_type}')
def then_inline_shape_type_is_shape_type(context, shape_type):
    expected_value = {
        'WD_INLINE_SHAPE.CHART':          WD_INLINE_SHAPE.CHART,
        'WD_INLINE_SHAPE.LINKED_PICTURE': WD_INLINE_SHAPE.LINKED_PICTURE,
        'WD_INLINE_SHAPE.PICTURE':        WD_INLINE_SHAPE.PICTURE,
        'WD_INLINE_SHAPE.SMART_ART':      WD_INLINE_SHAPE.SMART_ART,
    }[shape_type]
    inline_shape = context.inline_shape
    assert inline_shape.type == expected_value


@then('the document contains the inline picture')
def then_the_document_contains_the_inline_picture(context):
    document = context.document
    picture_shape = document.inline_shapes[0]
    blip = picture_shape._inline.graphic.graphicData.pic.blipFill.blip
    rId = blip.embed
    image_part = document._document_part.related_parts[rId]
    image_sha1 = hashlib.sha1(image_part.blob).hexdigest()
    expected_sha1 = '79769f1e202add2e963158b532e36c2c0f76a70c'
    assert image_sha1 == expected_sha1, (
        "image SHA1 doesn't match, expected %s, got %s" %
        (expected_sha1, image_sha1)
    )


@then('the length of the inline shape collection is 5')
def then_len_of_inline_shape_collection_is_5(context):
    inline_shapes = context.document.inline_shapes
    shape_count = len(inline_shapes)
    assert shape_count == 5, 'got %s' % shape_count